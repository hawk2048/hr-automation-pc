"""HR Automation - Resume Parsing Service"""
import re
import json
from typing import Optional
from datetime import datetime
import python_magic
from pypdf import PdfReader
from docx import Document


class ResumeParser:
    """Parse resumes from PDF/DOCX to structured JSON"""

    def __init__(self):
        # Common patterns
        self.email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        self.phone_pattern = r'(?:\+?86)?1[3-9]\d{9}'
        self.age_pattern = r'(?:年龄|age)[:：]?\s*(\d{1,2})'
        self.edu_pattern = r'(?:学历|学历层次|教育背景)[:：]?\s*([^\n]+)'
        self.exp_years_pattern = r'(?:工作年限|经验|工作经验)[:：]?\s*(\d+)\s*(?:年|years)'

    def parse(self, file_path: str) -> dict:
        """Parse resume file and return structured data"""
        mime = python_magic.from_file(file_path, mime=True)
        
        if 'pdf' in mime.lower():
            return self._parse_pdf(file_path)
        elif 'word' in mime.lower() or 'officedocument' in mime.lower():
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {mime}")

    def _parse_pdf(self, file_path: str) -> dict:
        """Parse PDF resume"""
        text = ""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        except Exception as e:
            raise ValueError(f"Failed to read PDF: {e}")
        
        return self._extract_info(text)

    def _parse_docx(self, file_path: str) -> dict:
        """Parse DOCX resume"""
        text = ""
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
        except Exception as e:
            raise ValueError(f"Failed to read DOCX: {e}")
        
        return self._extract_info(text)

    def _extract_info(self, text: str) -> dict:
        """Extract structured information from raw text"""
        
        # Basic info
        emails = re.findall(self.email_pattern, text)
        phones = re.findall(self.phone_pattern, text)
        
        # Find age
        age_match = re.search(self.age_pattern, text)
        age = int(age_match.group(1)) if age_match else None
        
        # Find education
        edu_match = re.search(self.edu_pattern, text)
        education = edu_match.group(1).strip() if edu_match else None
        
        # Find experience years
        exp_match = re.search(self.exp_years_pattern, text)
        exp_years = int(exp_match.group(1)) if exp_match else None
        
        # Name (usually at the top)
        lines = text.strip().split('\n')
        name = lines[0].strip() if lines else "Unknown"
        # Clean up name - remove common suffixes
        name = re.sub(r'[^\u4e00-\u9fa5a-zA-Z\s]', '', name).strip()
        
        # Work experience extraction (simplified)
        work_experience = self._extract_work_experience(text)
        
        # Education extraction
        education_list = self._extract_education(text)
        
        # Skills extraction
        skills = self._extract_skills(text)
        
        return {
            "name": name,
            "email": emails[0] if emails else None,
            "phone": phones[0] if phones else None,
            "age": age,
            "gender": self._extract_gender(text),
            "location": self._extract_location(text),
            "education": education_list,
            "work_experience": work_experience,
            "skills": skills,
            "experience_years": exp_years,
            "resume_text": text[:5000],  # Keep first 5000 chars
        }

    def _extract_work_experience(self, text: str) -> list:
        """Extract work experience entries"""
        experiences = []
        
        # Look for common patterns
        exp_keywords = ['工作经历', '工作经历', '职业经历', ' employment', ' experience']
        
        # Simple extraction - find date patterns and surrounding text
        date_pattern = r'(\d{4}[-/年]\d{1,2}?(?:[-/至]?)\d{4}[-/年]?\d{1,2}?)'
        
        matches = re.finditer(date_pattern, text)
        for match in matches:
            # Get context around date
            start = max(0, match.start() - 100)
            end = min(len(text), match.end() + 100)
            context = text[start:end]
            
            # Try to extract company and title
            lines = context.split('\n')
            exp_entry = {
                "duration": match.group(1),
                "company": lines[1].strip() if len(lines) > 1 else None,
                "title": lines[0].strip() if lines else None,
            }
            
            if exp_entry["title"] and len(exp_entry["title"]) < 50:
                experiences.append(exp_entry)
        
        return experiences[:5]  # Limit to 5 entries

    def _extract_education(self, text: str) -> list:
        """Extract education entries"""
        edu_list = []
        
        edu_keywords = ['教育背景', ' education', '学历']
        edu_levels = ['博士', '硕士', '本科', '大专', '高中', '中专']
        
        lines = text.split('\n')
        for i, line in enumerate(lines):
            for level in edu_levels:
                if level in line:
                    # Try to get university name from next line
                    university = lines[i+1].strip() if i+1 < len(lines) else None
                    if university and len(university) > 2:
                        edu_list.append({
                            "degree": level,
                            "university": university,
                            "major": self._extract_major(line),
                        })
                        break
        
        return edu_list[:3]  # Limit to 3 entries

    def _extract_major(self, text: str) -> Optional[str]:
        """Extract major from education line"""
        majors = ['计算机', '软件', '电子', '机械', '化学', '物理', '数学', '经济', '管理', '金融', '会计', '人力', '工商', '通信', '自动化']
        for major in majors:
            if major in text:
                return major
        return None

    def _extract_skills(self, text: str) -> dict:
        """Extract skills"""
        hard_skills = []
        soft_skills = []
        
        # Common tech skills
        tech_keywords = [
            'Python', 'Java', 'JavaScript', 'C++', 'C#', 'Go', 'Rust', 'PHP', 'Ruby', 'Swift', 'Kotlin',
            'React', 'Vue', 'Angular', 'Node.js', 'Django', 'Spring', 'Flask', 'FastAPI',
            'SQL', 'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'Elasticsearch',
            'AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes', 'Linux', 'Git',
            'Machine Learning', 'AI', 'Deep Learning', 'TensorFlow', 'PyTorch',
            'HTML', 'CSS', 'TypeScript', 'GraphQL', 'REST', 'API',
            '项目管理', '产品经理', 'UI', 'UX', 'Figma', 'Sketch',
        ]
        
        text_lower = text.lower()
        for skill in tech_keywords:
            if skill.lower() in text_lower:
                if skill in ['项目管理', '产品经理', '沟通', '领导', '团队协作']:
                    soft_skills.append({"name": skill, "level": "熟练"})
                else:
                    hard_skills.append({"name": skill, "level": "熟练"})
        
        # Deduplicate
        hard_skills = list({s["name"]: s for s in hard_skills}.values())
        soft_skills = list({s["name"]: s for s in soft_skills}.values())
        
        return {
            "hard_skills": hard_skills[:15],
            "soft_skills": soft_skills[:5],
        }

    def _extract_gender(self, text: str) -> Optional[str]:
        """Extract gender"""
        if '男' in text[:200]:
            return "male"
        elif '女' in text[:200]:
            return "female"
        return None

    def _extract_location(self, text: str) -> Optional[str]:
        """Extract location"""
        # Look for city names in first 500 chars
        cities = ['北京', '上海', '深圳', '广州', '杭州', '成都', '武汉', '南京', '西安', '苏州', '重庆', '天津']
        first_part = text[:500]
        for city in cities:
            if city in first_part:
                return city
        return None


# Global instance
resume_parser = ResumeParser()


def parse_resume(file_path: str) -> dict:
    """Parse resume file and return structured data"""
    return resume_parser.parse(file_path)


def parse_resume_text(text: str) -> dict:
    """Parse resume from raw text"""
    return resume_parser._extract_info(text)